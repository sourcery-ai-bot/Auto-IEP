#!/usr/bin/env python
# coding: utf-8

# ## Auto-IEP

# > Install required libraries

# In[1]:

# >Import required libraries

# In[2]:


import os
import sys
import re
import numpy as np
import pandas as pd
import docx2txt
from docx import Document
import requests
import json
import  tkinter as tk
from tkinter import filedialog, ttk
import threading

#Constants

API_KEY = "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"
API_ENDPOINT = "https://api.openai.com/v1/chat/completions"

# > Function definition for test selection
# In[4]:
# >Function definition for handling WCJ test reports
# In[6]:

def handle_WCJ_test_files(selected_file):
    #Convert student file to tables
    doc = Document(selected_file)
    tables = doc.tables

    #Turn tables into dataframe
    def table_to_df(table):
        data = []
        for row in table.rows:
            data.append([cell.text for cell in row.cells])
        df = pd.DataFrame(data)
        df.columns = df.iloc[0]
        df = df.drop(0)
        return df
    dataframes = [table_to_df(table) for table in tables]


    # Extract the text from the document
    text = [p.text for p in doc.paragraphs if p.text.strip() != ""]

    # Extract the score ranges and classifications
    data = []
    for line in text:
        # Search for lines that contain a score range and a classification
        match = re.search(r'(\d+ and [a-zA-Z]+|\d+ to \d+)\s*(.*)', line)
        if match:
            score_range = match.group(1)
            classification = match.group(2)
            data.append((score_range, classification))

    # Create a DataFrame for Score Ranges
    df_scoreranges = pd.DataFrame(data, columns=['Standard Score (SS) Range', 'WJIV Classification'])

    #Flag for having run the basic data transformer
    ran = False

    #Data transformer for basic info, name, dob etc.
    for dataframe in dataframes:
        if 'Name: ' in dataframe.columns[0] and ran == False:
            data = {
                'Name': [dataframe.columns[0].split(':')[1].strip()],
                'DOB': [dataframe.iloc[0, 0].split(':')[1].strip()],
                'Sex': [dataframe.iloc[2, 0].split(':')[1].strip()],
                'ID': [dataframe.iloc[2, 1].split('ID:')[1].strip()]
            }
            dataframenew = pd.DataFrame(data)
            ran = True
            
    #Saving real student name in a string so it can be placed into final response
    dataframes[0] = dataframenew
    REALSTUDENTNAME = dataframes[0].iloc[0,0].split(', ')[1].strip()

    #Function to convert score ranges to integers
    def convert_range_to_integers(score_range):
        if 'and above' in score_range:
            return [int(score_range.split(' ')[0]), float('inf')]
        elif 'and Below' in score_range:
            return [float('-inf'), int(score_range.split(' ')[0])]
        else:
            return [int(x) for x in score_range.split(' to ')]

    #Apply the function to the dataframe
    df_scoreranges['Score Range'] = df_scoreranges['Standard Score (SS) Range'].apply(convert_range_to_integers)

    #Adding a proficiency column for scores
    def get_proficiency(score):
        for index, row in df_scoreranges.iterrows():
           if not score == '' and row['Score Range'][0] <= int(score) <= row['Score Range'][1]:
                return row['WJIV Classification']
        return 'Unknown'
    proficiency_dfs = []
    for i, dataframe in enumerate(dataframes):
        if 'Cluster' in dataframes[i].columns:
            dataframes[i]['Proficiency'] = dataframes[i]['Current Scores'].apply(get_proficiency)
            proficiency_dfs.append(dataframes[i])
    allScores = pd.concat(proficiency_dfs)
    allScores = allScores[allScores['Proficiency'] != 'Unknown']
    allScores = allScores.reindex(axis = 0)

    #Creating a sentences list for our prompt, starting with test proficiencies
    student_name = "Placeholder"
    sentences = []
    for index, row in allScores.iterrows():
        subject_area = row['Cluster']
        proficiency = row['Proficiency']
        sentence = f"{student_name} has {proficiency} proficiency in {subject_area}."
        sentences.append(sentence)

    #QUALITATIVE OBS added to sentences
    observations_dfs = []
    for i, dataframe in enumerate(dataframes):
        if 'Woodcock-Johnson IV Tests of Achievement Form A and Extended Qualitative Observations' in dataframes[i].columns:
            observations_dfs.append(dataframes[i])
    observations = pd.concat(observations_dfs, ignore_index=True)
    observations = observations.iloc[:, 1]
    observations = observations.to_frame()
    observations = observations.reindex(axis = 0)
    for index, row in observations.iterrows():
        comment = row['Woodcock-Johnson IV Tests of Achievement Form A and Extended Qualitative Observations'] 
        sentence = f"In regards to {comment}"
        #Replaces student name FERPA
        sentence = sentence.replace(REALSTUDENTNAME, "Placeholder")
        sentences.append(sentence)

    #TEST OBS added to sentences
    observations_dfstest = []
    for i, dataframe in enumerate(dataframes):
        if 'Woodcock-Johnson IV Tests of Achievement Form A and Extended Test Session Observations' in dataframes[i].columns:
            observations_dfstest.append(dataframes[i])
    observations = pd.concat(observations_dfstest, ignore_index=True)
    observations = observations.iloc[:, 1]
    observations = observations.to_frame()
    observations = observations.reindex(axis = 0)
    for index, row in observations.iterrows():
        comment = row['Woodcock-Johnson IV Tests of Achievement Form A and Extended Test Session Observations'] 
        sentence = f"In regards to {comment}."
        #Replaces student name FERPA
        sentence = sentence.replace(REALSTUDENTNAME, "Placeholder")
        sentences.append(sentence)

    #Joining sentences for use in prompt
    dataForPrompt = ' '.join(map(str, sentences))

    #Tuning our input for best results
    promptTuning = """Write an academic assesment report based on the following data collected 
    from the Woodcock-Johnson IV Tests of Achievement, make it do not make it flowery and also understand this student 
    is in special education. 
    Be sure to highlight relative strengths, also do not write in all caps. 
    Avoid using run on sentences, try not to combine multiple subject areas into one sentence.
    Be certain to provide specific details. 
    Be cautious to avoid any kind of linguistic mistakes.
    Do not make special mention of referring to the student as Placeholder, treat it as a name.
    Data: """

    #Beginning of AI script
    #Openai interaction function
    def generate_chat_completion(messages, model="gpt-4", temperature=1, max_tokens=None):
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {API_KEY}",
        }

        data = {
            "model": model,
            "messages": messages,
            "temperature": temperature,
        }

        if max_tokens is not None:
            data["max_tokens"] = max_tokens

        response = requests.post(API_ENDPOINT, headers=headers, data=json.dumps(data))

        if response.status_code == 200:
            return response.json()["choices"][0]["message"]["content"]
        else:
            raise Exception(f"Error {response.status_code}: {response.text}")
    
    #The prompt we pass to gpt
    PROMPT = ( f"{promptTuning}{dataForPrompt}")
    
    #Define how gpt should act
    messages=[
        {"role": "system", "content": "You are a special education teacher"},
        {"role": "user", "content": (f"{PROMPT}")},
    ]
    #Call and get response
    response_text = generate_chat_completion(messages)
    #Replace placeholder with real student name
    #Due to FERPA considerations this may be commented out to keep the placeholder name for outside viewing
    response_text = response_text.replace("Placeholder", REALSTUDENTNAME)
    return (response_text)


# > GUI THINGS --------------------------------------------------------------------------
class AutoIEPGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("Auto-IEP")
        self.root.geometry("800x600")
        self.text_frame = None
        self.text_widget = None
        self.result_title = None
        self.save_button = None
        self.new_report_button = None
        self.progress = None
        self.scrollbar = None
        self.response_text = ""
        
        self.main_frame = tk.Frame(self.root)
        self.main_frame.pack(pady=20)
        
        label = tk.Label(self.main_frame, text="Supported tests are:")
        label.pack(pady=10)
        
        self.selected_test_var = tk.StringVar(root)
        test_dropdown = ttk.Combobox(self.main_frame, textvariable=self.selected_test_var, values=["1: WCJ"])
        test_dropdown.pack(pady=10)
        test_dropdown.current(0)
        
        process_button = tk.Button(self.main_frame, text="Select File and Process", command=self.main_program)
        process_button.pack(pady=20)
        
        self.progress = ttk.Progressbar(self.root, orient="horizontal", length=300, mode="determinate")
        self.spinner_label = tk.Label(self.root, text="Creating report :)")

    def on_api_call_complete(self, response_text):
        self.response_text = response_text
        self.progress.pack_forget()
        self.spinner_label.pack_forget()

        self.result_title = tk.Label(self.root, text="Auto IEP Report: ", font=("Arial", 12, "bold"))
        self.result_title.pack(pady=20)
        
        self.text_frame = tk.Frame(self.root)  # Change this line to make text_frame an instance variable
        self.text_frame.pack(pady=20, fill=tk.BOTH, expand=1)

        self.text_widget = tk.Text(self.text_frame, wrap=tk.WORD, font=("Arial", 12))
        self.text_widget.insert(tk.END, self.response_text)
        self.text_widget.pack(side=tk.LEFT, fill=tk.BOTH, expand=1)

        self.scrollbar = tk.Scrollbar(self.text_frame, command=self.text_widget.yview)
        self.scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        self.text_widget.config(yscrollcommand=self.scrollbar.set)

        self.save_button = tk.Button(self.root, text="Save Results", command=self.save_results)
        self.save_button.pack(pady=20)

        self.new_report_button = tk.Button(self.root, text="Generate New Report", command=self.generate_new_report)
        self.new_report_button.pack(pady=20)

    def api_call_thread(self, selected_file):
        non_thread_response = handle_WCJ_test_files(selected_file)
        self.root.after(10, lambda: self.on_api_call_complete(non_thread_response))

    def delayed_file_selection(self):
        selected_file = self.file_select_WCJ()
        if selected_file:
            self.main_frame.pack_forget()
            self.show_progress_view()
            self.update_progress()
            threading.Thread(target=self.api_call_thread, args=(selected_file,)).start()

    def update_progress(self):
        current_value = self.progress["value"]
        new_value = current_value + 0.2
        self.progress["value"] = new_value
        self.root.after(100, self.update_progress)

    def main_program(self):
        selected_test = self.selected_test_var.get()
        if selected_test == "1: WCJ":
            self.root.after(100, self.delayed_file_selection)

    def show_progress_view(self):
        self.progress.pack(pady=20)
        self.spinner_label.pack(pady=20)

    def save_results(self):
        file_name = filedialog.asksaveasfilename(title="Save Results", filetypes=[("Text files", "*.txt"), ("All files", "*.*")], defaultextension=".txt")
        if file_name:
            with open(file_name, 'w') as file:
                file.write(self.response_text)
            tk.messagebox.showinfo("Success", "Results saved successfully!")

    def generate_new_report(self):
        if self.text_widget:
            self.text_widget.pack_forget()
        if self.text_frame:
            self.text_frame.pack_forget()
            
        # Hide other widgets
        if self.result_title:
            self.result_title.pack_forget()
        if self.save_button:
            self.save_button.pack_forget()
        if self.new_report_button:
            self.new_report_button.pack_forget()
        if self.scrollbar:
            self.scrollbar.pack_forget()
        if self.progress:
            self.progress.pack_forget()
        if self.spinner_label:
            self.spinner_label.pack_forget()
            
        self.response_text = ""
        self.progress["value"] = 0
        self.main_frame.pack(pady=20)

    def file_select_WCJ(self):
        self.root.withdraw()  # hide main window
        file_path = filedialog.askopenfilename(title = "Select a WCJ File")
        self.root.deiconify()  # show main window
        return file_path

# Execution:
root = tk.Tk()
app = AutoIEPGUI(root)
root.mainloop()
# In[ ]:




